#!/usr/bin/env python3
"""Generate or edit a .docx file from a JSON spec.

Usage:
    # Generate from scratch
    python generate_docx.py <spec.json> <output.docx>

    # Edit an existing file: apply replacements then append blocks
    python generate_docx.py <spec.json> <output.docx> --base <existing.docx>

Spec schema: see ../references/spec_reference.md
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    sys.stderr.write(
        "error: python-docx is not installed. Run: pip install python-docx\n"
    )
    sys.exit(1)


ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# --- Helpers --------------------------------------------------------------

def _parse_color(value):
    """Accept 'RRGGBB', '#RRGGBB', or [r,g,b] and return RGBColor."""
    if value is None:
        return None
    if isinstance(value, (list, tuple)) and len(value) == 3:
        return RGBColor(int(value[0]), int(value[1]), int(value[2]))
    s = str(value).lstrip("#")
    if len(s) != 6:
        raise ValueError(f"Invalid color: {value!r} (expected 'RRGGBB')")
    return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))


def _apply_run_format(run, fmt):
    if "bold" in fmt:
        run.bold = bool(fmt["bold"])
    if "italic" in fmt:
        run.italic = bool(fmt["italic"])
    if "underline" in fmt:
        run.underline = bool(fmt["underline"])
    if "size" in fmt:
        run.font.size = Pt(fmt["size"])
    if "color" in fmt:
        color = _parse_color(fmt["color"])
        if color is not None:
            run.font.color.rgb = color
    if "font" in fmt:
        run.font.name = fmt["font"]


def _add_runs(paragraph, runs_spec):
    """A run is either a string or {text, bold?, italic?, underline?, size?, color?, font?, break?}."""
    for item in runs_spec:
        if isinstance(item, str):
            paragraph.add_run(item)
            continue
        text = item.get("text", "")
        run = paragraph.add_run(text)
        _apply_run_format(run, item)
        if item.get("break") == "line":
            run.add_break(WD_BREAK.LINE)
        elif item.get("break") == "page":
            run.add_break(WD_BREAK.PAGE)


def _set_paragraph_text(paragraph, spec):
    """Populate a paragraph from {text} or {runs}. Apply alignment."""
    runs_spec = spec.get("runs")
    if runs_spec:
        _add_runs(paragraph, runs_spec)
    else:
        text = spec.get("text", "")
        if text:
            run = paragraph.add_run(text)
            # Allow paragraph-level run formatting shortcuts
            _apply_run_format(run, spec)
    align = spec.get("align")
    if align:
        if align not in ALIGN_MAP:
            raise ValueError(
                f"Unknown align: {align!r}. Use one of: {', '.join(ALIGN_MAP)}"
            )
        paragraph.alignment = ALIGN_MAP[align]


def _add_horizontal_rule(paragraph):
    """Insert a bottom border on the paragraph to render as a horizontal line."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _ensure_list_styles(doc):
    """python-docx ships with 'List Bullet' and 'List Number' styles in default templates."""
    return doc.styles


# --- Block handlers -------------------------------------------------------

def _block_heading(doc, spec, _base_dir):
    level = int(spec.get("level", 1))
    if level < 0 or level > 9:
        raise ValueError(f"Invalid heading level {level}; must be 0–9")
    text = spec.get("text", "")
    p = doc.add_heading(text, level=level)
    align = spec.get("align")
    if align:
        p.alignment = ALIGN_MAP[align]


def _block_paragraph(doc, spec, _base_dir):
    style = spec.get("style")  # e.g. "Normal", "Intense Quote"
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    _set_paragraph_text(p, spec)


def _block_bullets(doc, spec, _base_dir):
    items = spec.get("items", [])
    style = spec.get("style", "List Bullet")
    for item in items:
        p = doc.add_paragraph(style=style)
        if isinstance(item, dict):
            _set_paragraph_text(p, item)
        else:
            p.add_run(str(item))


def _block_numbered(doc, spec, _base_dir):
    items = spec.get("items", [])
    style = spec.get("style", "List Number")
    for item in items:
        p = doc.add_paragraph(style=style)
        if isinstance(item, dict):
            _set_paragraph_text(p, item)
        else:
            p.add_run(str(item))


def _block_table(doc, spec, _base_dir):
    headers = spec.get("headers", [])
    rows = spec.get("rows", [])
    if not headers and not rows:
        return
    n_cols = len(headers) if headers else len(rows[0])
    n_rows = len(rows) + (1 if headers else 0)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = spec.get("table_style", "Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    r = 0
    if headers:
        for c, h in enumerate(headers):
            cell = table.cell(0, c)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(h))
            run.bold = True
        r = 1
    for row in rows:
        for c, val in enumerate(row[:n_cols]):
            cell = table.cell(r, c)
            cell.text = ""
            cell.paragraphs[0].add_run(str(val))
        r += 1


def _block_image(doc, spec, base_dir):
    img_path = Path(spec["image_path"])
    if not img_path.is_absolute():
        img_path = base_dir / img_path
    if not img_path.is_file():
        raise FileNotFoundError(f"Image not found: {img_path}")
    width = spec.get("width_inches")
    kwargs = {"width": Inches(width)} if width else {}
    p = doc.add_paragraph()
    p.alignment = ALIGN_MAP[spec.get("align", "center")]
    p.add_run().add_picture(str(img_path), **kwargs)
    caption = spec.get("caption")
    if caption:
        cap = doc.add_paragraph(style="Caption") if _has_style(doc, "Caption") else doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True


def _has_style(doc, name):
    try:
        doc.styles[name]
        return True
    except KeyError:
        return False


def _block_quote(doc, spec, _base_dir):
    style = "Intense Quote" if _has_style(doc, "Intense Quote") else "Quote" if _has_style(doc, "Quote") else None
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    text = spec.get("text", "")
    run = p.add_run(text)
    if not style:
        run.italic = True
    attribution = spec.get("attribution")
    if attribution:
        ap = doc.add_paragraph()
        ap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ar = ap.add_run(f"— {attribution}")
        ar.italic = True


def _block_code(doc, spec, _base_dir):
    text = spec.get("text", "")
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)


def _block_page_break(doc, _spec, _base_dir):
    doc.add_page_break()


def _block_divider(doc, _spec, _base_dir):
    p = doc.add_paragraph()
    _add_horizontal_rule(p)


HANDLERS = {
    "heading": _block_heading,
    "paragraph": _block_paragraph,
    "bullets": _block_bullets,
    "numbered": _block_numbered,
    "table": _block_table,
    "image": _block_image,
    "quote": _block_quote,
    "code": _block_code,
    "page_break": _block_page_break,
    "divider": _block_divider,
}


# --- Find/replace ---------------------------------------------------------

def _replace_in_paragraph(paragraph, find, replace):
    if find not in paragraph.text:
        return 0
    full = paragraph.text
    count = full.count(find)
    new_text = full.replace(find, replace)
    runs = paragraph.runs
    if not runs:
        return 0
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""
    return count


def _apply_replacements(doc, replacements):
    total = 0
    for rep in replacements:
        find = rep.get("find")
        replace = rep.get("replace", "")
        if find is None or find == "":
            continue
        for paragraph in doc.paragraphs:
            total += _replace_in_paragraph(paragraph, find, replace)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        total += _replace_in_paragraph(paragraph, find, replace)
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                total += _replace_in_paragraph(paragraph, find, replace)
            for paragraph in section.footer.paragraphs:
                total += _replace_in_paragraph(paragraph, find, replace)
    return total


# --- Build ----------------------------------------------------------------

def build(spec, output, base_dir, base_path=None):
    if base_path:
        if not base_path.is_file():
            raise FileNotFoundError(f"Base document not found: {base_path}")
        doc = Document(str(base_path))
    else:
        doc = Document()

    if spec.get("title"):
        doc.core_properties.title = spec["title"]
    if spec.get("author"):
        doc.core_properties.author = spec["author"]

    replaced = 0
    if base_path:
        replaced = _apply_replacements(doc, spec.get("replacements", []))

    blocks = spec.get("blocks", [])
    for i, block in enumerate(blocks):
        btype = block.get("type")
        if btype is None:
            raise ValueError(f"Block {i} missing 'type'")
        handler = HANDLERS.get(btype)
        if handler is None:
            raise ValueError(
                f"Unknown block type at index {i}: {btype!r}. "
                f"Valid types: {', '.join(HANDLERS)}"
            )
        handler(doc, block, base_dir)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return len(blocks), replaced


def main(argv):
    args = argv[1:]
    base_path = None
    rest = []
    i = 0
    while i < len(args):
        a = args[i]
        if a == "--base":
            if i + 1 >= len(args):
                sys.stderr.write("error: --base requires a path\n")
                return 2
            base_path = Path(args[i + 1]).resolve()
            i += 2
        else:
            rest.append(a)
            i += 1

    if len(rest) != 2:
        sys.stderr.write(
            "usage: generate_docx.py <spec.json> <output.docx> [--base <existing.docx>]\n"
        )
        return 2

    spec_path = Path(rest[0]).resolve()
    out_path = Path(rest[1]).resolve()
    if out_path.suffix.lower() != ".docx":
        sys.stderr.write("error: output path must end in .docx\n")
        return 2
    if not spec_path.is_file():
        sys.stderr.write(f"error: spec not found: {spec_path}\n")
        return 2

    spec = json.loads(spec_path.read_text())
    n_blocks, n_replaced = build(spec, out_path, spec_path.parent, base_path)
    if base_path:
        print(
            f"Edited {base_path.name}: {n_replaced} replacement(s), "
            f"{n_blocks} block(s) appended → {out_path}"
        )
    else:
        print(f"Wrote {n_blocks} block(s) → {out_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
